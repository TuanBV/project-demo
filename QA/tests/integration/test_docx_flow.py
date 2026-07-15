"""Spec section 23 'Luồng DOCX': upload -> dry-run -> preview -> real import -> reimport dedup."""

from __future__ import annotations

import io

import docx
from fastapi.testclient import TestClient


def _build_fixture_docx() -> bytes:
    document = docx.Document()
    document.add_paragraph("PHAN I - JAVA CORE")
    document.add_paragraph("")
    document.add_paragraph("Cau 1. JVM la gi?")
    document.add_paragraph("")
    document.add_paragraph("Tra loi: JVM la may ao thuc thi Java bytecode.")
    document.add_paragraph("")
    document.add_paragraph("Cau 2. Bytecode la gi?")
    document.add_paragraph("")
    document.add_paragraph("Tra loi: Bytecode la ma trung gian duoc JVM thuc thi.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _upload(client: TestClient, dry_run: bool, duplicate_strategy: str = "SKIP"):
    files = {
        "file": (
            "fixture.docx",
            io.BytesIO(_build_fixture_docx()),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }
    data = {"dry_run": str(dry_run), "duplicate_strategy": duplicate_strategy}
    return client.post("/api/admin/import/docx", files=files, data=data)


def test_docx_dry_run_preview(client: TestClient) -> None:
    response = _upload(client, dry_run=True)
    assert response.status_code == 200
    body = response.json()
    assert body["dry_run"] is True
    assert body["job_id"] is None
    assert body["summary"]["questions_detected"] == 2
    assert body["summary"]["questions_created"] == 0

    categories = client.get("/api/categories").json()
    assert categories == []


def test_docx_real_import_creates_category_and_questions(client: TestClient) -> None:
    response = _upload(client, dry_run=False)
    body = response.json()
    assert body["summary"]["categories_created"] == 1
    assert body["summary"]["questions_created"] == 2
    assert body["summary"]["questions_needs_review"] == 2
    assert all(item["status"] == "NEEDS_REVIEW" for item in body["items"])

    categories = client.get("/api/categories").json()
    assert len(categories) == 1
    assert categories[0]["name"] == "JAVA CORE"

    admin_questions = client.get("/api/admin/questions").json()
    assert len(admin_questions) == 2


def test_docx_import_creates_auto_generated_distractors_pending_review(
    client: TestClient,
) -> None:
    """DOCX only has question+correct-answer (spec 9): importer must add 3 distractor
    suggestions marked auto_generated, and keep the question inactive until reviewed."""
    _upload(client, dry_run=False)
    admin_questions = client.get("/api/admin/questions").json()
    question_id = admin_questions[0]["id"]

    detail = client.get(f"/api/admin/questions/{question_id}").json()
    assert detail["needs_review"] is True
    assert detail["active"] is False
    assert len(detail["options"]) == 4
    correct = [o for o in detail["options"] if o["is_correct"]]
    distractors = [o for o in detail["options"] if not o["is_correct"]]
    assert len(correct) == 1
    assert len(distractors) == 3
    assert all(o["auto_generated"] for o in distractors)
    assert not correct[0]["auto_generated"]

    # Not active -> must not appear in the study flow yet.
    session = client.post("/api/study-sessions", json={"mode": "RANDOM"}).json()
    next_response = client.post(f"/api/study-sessions/{session['id']}/next")
    assert next_response.status_code == 404

    # Admin reviews and activates -> now importable/studyable.
    update = client.put(
        f"/api/admin/questions/{question_id}",
        json={"active": True, "needs_review": False},
    )
    assert update.status_code == 200
    validation = client.post(f"/api/admin/questions/{question_id}/validate").json()
    assert validation["status"] == "VALID"

    next_response2 = client.post(f"/api/study-sessions/{session['id']}/next")
    assert next_response2.status_code == 200


def test_docx_reimport_is_deduplicated(client: TestClient) -> None:
    _upload(client, dry_run=False)
    second = _upload(client, dry_run=False, duplicate_strategy="SKIP")
    body = second.json()
    assert body["summary"]["questions_created"] == 0
    assert body["summary"]["questions_skipped"] == 2

    admin_questions = client.get("/api/admin/questions").json()
    assert len(admin_questions) == 2
