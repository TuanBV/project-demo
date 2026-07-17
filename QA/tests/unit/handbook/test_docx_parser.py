"""Parser tests for app/handbook/docx_parser.py. Split into:
- tests against the REAL source DOCX (scripts/data/so_tay_on_tap_de_doc_noi_bat.docx) -- the
  only data source allowed for this feature.
- tests against small synthetic DOCX files built with python-docx, to exercise error paths
  without touching the real source file.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from app.handbook.docx_parser import DocxStructureError, parse_docx
from app.handbook.store import DOCX_PATH

pytestmark = pytest.mark.skipif(
    not DOCX_PATH.exists(), reason="scripts/data/so_tay_on_tap_de_doc_noi_bat.docx missing"
)

_TOPIC_COUNT = 12
_QUESTION_COUNT = 240


def _ensure_style(document: Document, name: str) -> None:
    if name not in document.styles:
        document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _add_topic_heading(document: Document, order: int, name: str) -> None:
    document.add_heading(f"CHỦ ĐỀ {order:02d} — {name}", level=1)


def _add_terms_table(document: Document, terms: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1 + len(terms), cols=2)
    table.rows[0].cells[0].text = "Thuật ngữ / Khái niệm"
    table.rows[0].cells[1].text = "Giải thích dễ nhớ"
    for i, (term, definition) in enumerate(terms, start=1):
        table.rows[i].cells[0].text = term
        table.rows[i].cells[1].text = definition


def _add_card(
    document: Document, number: int, question: str, core: str, deep: str | None = None
) -> None:
    _ensure_style(document, "Memory Card")
    document.add_paragraph(f"CÂU {number}  |  {question}", style="Memory Card")
    document.add_paragraph(f"TRẢ LỜI CỐT LÕI:  {core}", style="Memory Card")
    if deep is not None:
        document.add_paragraph(f"GIẢI THÍCH:  {deep}", style="Memory Card")


def _build_minimal_valid_docx(
    path: Path, topic_count: int = _TOPIC_COUNT, questions_per_topic: int = 1
) -> None:
    document = Document()
    _ensure_style(document, "Code Block")
    number = 1
    for topic_index in range(1, topic_count + 1):
        _add_topic_heading(document, topic_index, f"Chủ đề số {topic_index}")
        document.add_paragraph(f"Phạm vi câu hỏi: {number:03d}–{number:03d}")
        document.add_paragraph(f"Mục tiêu: Mục tiêu của chủ đề {topic_index}")
        _add_terms_table(document, [("Term", "Definition")])
        document.add_paragraph("Một sai lầm thường gặp.", style="List Bullet")
        document.add_paragraph("[Text]\nVí dụ cốt lõi", style="Code Block")
        for _ in range(questions_per_topic):
            _add_card(
                document,
                number,
                f"Câu hỏi số {number}?",
                f"Đáp án số {number}",
                f"Giải thích số {number}",
            )
            number += 1
    document.save(str(path))


class TestRealDocx:
    def test_parses_exactly_12_topics(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        assert len(parsed.topics) == _TOPIC_COUNT

    def test_parses_exactly_240_questions(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        assert sum(t.question_count for t in parsed.topics) == _QUESTION_COUNT

    def test_question_numbers_are_globally_unique(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        numbers = [q.number for t in parsed.topics for q in t.questions]
        assert len(numbers) == len(set(numbers))

    def test_every_topic_has_goal_and_range(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        for topic in parsed.topics:
            assert topic.goal.strip()
            assert topic.question_range.strip()

    def test_every_topic_has_terms_and_mistakes(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        for topic in parsed.topics:
            assert len(topic.terms) > 0
            assert len(topic.common_mistakes) > 0

    def test_code_block_placeholder_is_stripped(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        for topic in parsed.topics:
            for example in topic.core_examples:
                assert not example.startswith("[Text]")

    def test_some_questions_have_no_explanation(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        all_questions = [q for t in parsed.topics for q in t.questions]
        assert any(not q.explanation for q in all_questions)
        assert any(q.explanation for q in all_questions)
        assert all(q.question.strip() and q.core.strip() for q in all_questions)

    def test_topic_slugs_are_unique_and_expected(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        slugs = [t.slug for t in parsed.topics]
        assert len(slugs) == len(set(slugs))
        assert slugs == [
            "ai-coding",
            "du-an-production",
            "java-core",
            "oop-solid",
            "spring-boot",
            "git",
            "thuat-toan",
            "rest-api",
            "sql-database",
            "jvm-performance",
            "microservices",
            "cau-hoi-ntd",
        ]

    def test_appendix_glossary_is_parsed(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        assert len(parsed.glossary) > 0
        assert all(t.term.strip() and t.definition.strip() for t in parsed.glossary)

    def test_preserves_vietnamese_unicode(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        question_one = next(q for t in parsed.topics for q in t.questions if q.number == 1)
        assert (
            "ế" in question_one.question
            or "ự" in question_one.question
            or "ậ" in question_one.question
            or "ệ" in question_one.core
        )

    def test_does_not_read_java_python_mc_json(self, monkeypatch: pytest.MonkeyPatch) -> None:
        import pathlib

        original_read_text = pathlib.Path.read_text

        def _guarded_read_text(self: pathlib.Path, *args: object, **kwargs: object) -> str:
            assert "java_python_mc" not in str(self), f"Parser must not read {self}"
            return original_read_text(self, *args, **kwargs)

        monkeypatch.setattr(pathlib.Path, "read_text", _guarded_read_text)
        parse_docx(DOCX_PATH)


class TestSyntheticDocxErrorPaths:
    def test_valid_minimal_docx_parses(self, tmp_path: Path) -> None:
        path = tmp_path / "valid.docx"
        _build_minimal_valid_docx(path, topic_count=_TOPIC_COUNT, questions_per_topic=2)
        parsed = parse_docx(path)
        assert len(parsed.topics) == _TOPIC_COUNT
        assert sum(t.question_count for t in parsed.topics) == _TOPIC_COUNT * 2
        assert parsed.topics[0].terms == [
            type(parsed.topics[0].terms[0])(term="Term", definition="Definition")
        ]
        assert parsed.topics[0].common_mistakes == ["Một sai lầm thường gặp."]
        assert parsed.topics[0].core_examples == ["Ví dụ cốt lõi"]

    def test_card_without_explanation_parses_with_empty_explanation(self, tmp_path: Path) -> None:
        document = Document()
        _ensure_style(document, "Code Block")
        for topic_index in range(1, _TOPIC_COUNT + 1):
            _add_topic_heading(document, topic_index, f"Chủ đề số {topic_index}")
            document.add_paragraph(f"Phạm vi câu hỏi: {topic_index:03d}")
            document.add_paragraph(f"Mục tiêu: Mục tiêu {topic_index}")
            _add_terms_table(document, [("T", "D")])
            _add_card(document, topic_index, f"Câu {topic_index}?", f"Đáp án {topic_index}")
        path = tmp_path / "no_explanation.docx"
        document.save(str(path))
        parsed = parse_docx(path)
        all_questions = [q for t in parsed.topics for q in t.questions]
        assert all(q.explanation == "" for q in all_questions)
        assert all(q.core for q in all_questions)

    def test_missing_topic_raises(self, tmp_path: Path) -> None:
        path = tmp_path / "missing_topic.docx"
        _build_minimal_valid_docx(path, topic_count=_TOPIC_COUNT - 1, questions_per_topic=1)
        with pytest.raises(DocxStructureError, match="chủ đề"):
            parse_docx(path)

    def test_too_many_topics_raises(self, tmp_path: Path) -> None:
        path = tmp_path / "extra_topic.docx"
        _build_minimal_valid_docx(path, topic_count=_TOPIC_COUNT + 1, questions_per_topic=1)
        with pytest.raises(DocxStructureError):
            parse_docx(path)

    def test_duplicate_question_number_raises(self, tmp_path: Path) -> None:
        document = Document()
        _add_topic_heading(document, 1, "A")
        _add_card(document, 1, "Q1?", "A1", "E1")
        _add_card(document, 1, "Q1 again?", "A1b", "E1b")
        for topic_index in range(2, _TOPIC_COUNT + 1):
            _add_topic_heading(document, topic_index, f"Topic {topic_index}")
            _add_card(document, topic_index + 1, f"Q{topic_index}?", "A", "E")
        path = tmp_path / "dup.docx"
        document.save(str(path))
        with pytest.raises(DocxStructureError, match="trùng"):
            parse_docx(path)

    def test_missing_core_raises(self, tmp_path: Path) -> None:
        document = Document()
        _add_topic_heading(document, 1, "A")
        _ensure_style(document, "Memory Card")
        document.add_paragraph("CÂU 1  |  Question only, no core line", style="Memory Card")
        for topic_index in range(2, _TOPIC_COUNT + 1):
            _add_topic_heading(document, topic_index, f"Topic {topic_index}")
            _add_card(document, topic_index, f"Q{topic_index}?", "A", "E")
        path = tmp_path / "missing_core.docx"
        document.save(str(path))
        with pytest.raises(DocxStructureError):
            parse_docx(path)

    def test_missing_file_raises(self, tmp_path: Path) -> None:
        with pytest.raises(DocxStructureError, match="Không tìm thấy"):
            parse_docx(tmp_path / "does_not_exist.docx")
