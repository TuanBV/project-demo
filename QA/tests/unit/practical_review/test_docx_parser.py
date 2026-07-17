"""Parser tests for app/practical_review/docx_parser.py. Split into:
- tests against the REAL source DOCX
  (scripts/data/so_tay_on_tap_sap_xep_theo_chu_de_uu_tien.docx) -- the only data source
  allowed for this feature.
- tests against small synthetic DOCX files built with python-docx, to exercise error paths
  (missing field, duplicate number, wrong topic count) without touching the real source file.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from app.practical_review.docx_parser import DocxStructureError, parse_docx
from app.practical_review.store import DOCX_PATH

pytestmark = pytest.mark.skipif(
    not DOCX_PATH.exists(),
    reason="scripts/data/so_tay_on_tap_sap_xep_theo_chu_de_uu_tien.docx missing",
)

_TOPIC_COUNT = 11
_QUESTION_COUNT = 207


def _ensure_memory_card_style(document: Document) -> None:
    if "Memory Card" not in document.styles:
        document.styles.add_style("Memory Card", WD_STYLE_TYPE.PARAGRAPH)


def _add_topic_heading(document: Document, order: int, name: str) -> None:
    document.add_heading(f"CHỦ ĐỀ {order:02d} — {name}", level=1)


def _add_card(
    document: Document, number: int, question: str, core: str, deep: str | None = None
) -> None:
    _ensure_memory_card_style(document)
    document.add_paragraph(f"{number}. {question}", style="Memory Card")
    document.add_paragraph(f"Cốt lõi: {core}", style="Memory Card")
    if deep is not None:
        document.add_paragraph(f"Hiểu sâu: {deep}", style="Memory Card")


def _build_minimal_valid_docx(
    path: Path, topic_count: int = _TOPIC_COUNT, questions_per_topic: int = 1
) -> None:
    document = Document()
    number = 1
    for topic_index in range(1, topic_count + 1):
        _add_topic_heading(document, topic_index, f"Chủ đề số {topic_index}")
        for _ in range(questions_per_topic):
            _add_card(
                document,
                number,
                f"Câu hỏi số {number}?",
                f"Đáp án số {number}",
                f"Giải thích số {number}",
            )
            number += 1
    document.save(str(path))


class TestRealDocx:
    def test_parses_exactly_11_topics(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        assert len(parsed.topics) == _TOPIC_COUNT

    def test_parses_exactly_207_questions(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        assert len(parsed.questions) == _QUESTION_COUNT

    def test_question_numbers_are_globally_unique(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        numbers = [q.number for q in parsed.questions]
        assert len(numbers) == len(set(numbers))

    def test_question_numbers_within_original_1_to_240_range(self) -> None:
        # The handbook is a priority-reordered SUBSET of a larger original bank -- numbering
        # is not contiguous (whole ranges like 141-160 are intentionally dropped).
        parsed = parse_docx(DOCX_PATH)
        numbers = [q.number for q in parsed.questions]
        assert min(numbers) >= 1
        assert max(numbers) <= 240

    def test_every_question_has_non_empty_question_and_answer(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        for question in parsed.questions:
            assert question.question.strip()
            assert question.answer.strip()

    def test_some_questions_have_no_deep_explanation(self) -> None:
        # Not every "Memory Card" has a "Hiểu sâu:" line -- explanation is allowed to be "".
        parsed = parse_docx(DOCX_PATH)
        assert any(not q.explanation for q in parsed.questions)
        assert any(q.explanation for q in parsed.questions)

    def test_preserves_vietnamese_unicode(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        question_one = next(q for q in parsed.questions if q.number == 1)
        # Real Vietnamese diacritics must round-trip exactly, not mojibake.
        assert (
            "ế" in question_one.question
            or "ự" in question_one.question
            or "ậ" in question_one.question
            or "ệ" in question_one.answer
        )

    def test_topic_slugs_are_unique_and_expected(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        slugs = [t.slug for t in parsed.topics]
        assert len(slugs) == len(set(slugs))
        assert slugs == [
            "ai-coding",
            "du-an-production",
            "java-core",
            "oop-solid",
            "spring-boot",
            "git",
            "thuat-toan",
            "rest-api",
            "sql-database",
            "python-core",
            "python-backend",
        ]

    def test_topic_question_counts_sum_to_total(self) -> None:
        parsed = parse_docx(DOCX_PATH)
        assert sum(t.question_count for t in parsed.topics) == _QUESTION_COUNT

    def test_does_not_read_java_python_mc_json(self, monkeypatch: pytest.MonkeyPatch) -> None:
        """Guard: the parser must only touch the DOCX file, never
        scripts/data/java_python_mc/*.json."""
        import pathlib

        original_read_text = pathlib.Path.read_text

        def _guarded_read_text(self: pathlib.Path, *args: object, **kwargs: object) -> str:
            assert "java_python_mc" not in str(self), f"Parser must not read {self}"
            return original_read_text(self, *args, **kwargs)

        monkeypatch.setattr(pathlib.Path, "read_text", _guarded_read_text)
        parse_docx(DOCX_PATH)


class TestSyntheticDocxErrorPaths:
    def test_valid_minimal_docx_parses(self, tmp_path: Path) -> None:
        path = tmp_path / "valid.docx"
        _build_minimal_valid_docx(path, topic_count=_TOPIC_COUNT, questions_per_topic=2)
        parsed = parse_docx(path)
        assert len(parsed.topics) == _TOPIC_COUNT
        assert len(parsed.questions) == _TOPIC_COUNT * 2

    def test_card_without_deep_line_parses_with_empty_explanation(self, tmp_path: Path) -> None:
        document = Document()
        for topic_index in range(1, _TOPIC_COUNT + 1):
            _add_topic_heading(document, topic_index, f"Chủ đề số {topic_index}")
            _add_card(document, topic_index, f"Câu {topic_index}?", f"Đáp án {topic_index}")
        path = tmp_path / "no_deep.docx"
        document.save(str(path))
        parsed = parse_docx(path)
        assert all(q.explanation == "" for q in parsed.questions)
        assert all(q.answer for q in parsed.questions)

    def test_missing_topic_raises(self, tmp_path: Path) -> None:
        path = tmp_path / "missing_topic.docx"
        _build_minimal_valid_docx(path, topic_count=_TOPIC_COUNT - 1, questions_per_topic=1)
        with pytest.raises(DocxStructureError, match="chủ đề"):
            parse_docx(path)

    def test_too_many_topics_raises(self, tmp_path: Path) -> None:
        path = tmp_path / "extra_topic.docx"
        _build_minimal_valid_docx(path, topic_count=_TOPIC_COUNT + 1, questions_per_topic=1)
        with pytest.raises(DocxStructureError):
            parse_docx(path)

    def test_duplicate_question_number_raises(self, tmp_path: Path) -> None:
        document = Document()
        _add_topic_heading(document, 1, "A")
        _add_card(document, 1, "Q1?", "A1", "E1")
        _add_card(document, 1, "Q1 again?", "A1b", "E1b")
        for topic_index in range(2, _TOPIC_COUNT + 1):
            _add_topic_heading(document, topic_index, f"Topic {topic_index}")
            _add_card(document, topic_index + 1, f"Q{topic_index}?", "A", "E")
        path = tmp_path / "dup.docx"
        document.save(str(path))
        with pytest.raises(DocxStructureError, match="trùng"):
            parse_docx(path)

    def test_missing_core_raises(self, tmp_path: Path) -> None:
        document = Document()
        _add_topic_heading(document, 1, "A")
        _ensure_memory_card_style(document)
        document.add_paragraph("1. Question only, no Cốt lõi line", style="Memory Card")
        for topic_index in range(2, _TOPIC_COUNT + 1):
            _add_topic_heading(document, topic_index, f"Topic {topic_index}")
            _add_card(document, topic_index, f"Q{topic_index}?", "A", "E")
        path = tmp_path / "missing_core.docx"
        document.save(str(path))
        with pytest.raises(DocxStructureError):
            parse_docx(path)

    def test_card_before_any_topic_heading_raises(self, tmp_path: Path) -> None:
        document = Document()
        _ensure_memory_card_style(document)
        document.add_paragraph("1. Orphan question", style="Memory Card")
        document.add_paragraph("Cốt lõi: orphan answer", style="Memory Card")
        path = tmp_path / "orphan.docx"
        document.save(str(path))
        with pytest.raises(DocxStructureError):
            parse_docx(path)

    def test_missing_file_raises(self, tmp_path: Path) -> None:
        with pytest.raises(DocxStructureError, match="Không tìm thấy"):
            parse_docx(tmp_path / "does_not_exist.docx")
