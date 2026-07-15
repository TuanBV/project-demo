"""DocxTextExtractor: DOCX -> plain text, preserving paragraph order and blank lines.

Only responsible for extraction; question-recognition logic lives in the shared parsers
(interview_text_parser.py / structured_text_parser.py) so DOCX and pasted text never
duplicate business logic (spec section 5.1/5.2).
"""

from __future__ import annotations

import io

import docx

from app.core.exceptions import MalformedDocumentError


class DocxTextExtractor:
    def extract(self, source: bytes | io.BytesIO) -> str:
        buffer = io.BytesIO(source) if isinstance(source, bytes | bytearray) else source
        try:
            document = docx.Document(buffer)
        except Exception as exc:  # python-docx raises various exceptions for bad files
            raise MalformedDocumentError(f"Không thể đọc file DOCX: {exc}") from exc

        lines: list[str] = []
        for paragraph in document.paragraphs:
            lines.append(paragraph.text.rstrip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("\t".join(cells))

        text = "\n".join(lines).strip()
        if not text:
            raise MalformedDocumentError("File DOCX không chứa nội dung văn bản.")
        return text


class PlainTextExtractor:
    def extract(self, source: str) -> str:
        text = source.replace("\r\n", "\n").replace("\r", "\n")
        return text.strip("﻿").strip()
