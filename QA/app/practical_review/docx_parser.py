"""Parses scripts/data/so_tay_on_tap_sap_xep_theo_chu_de_uu_tien.docx -- the ONLY data source
for the practical-review area. Must never import or read anything under
scripts/data/java_python_mc/ or scripts/data/extended_topics/.

Expected structure (verified against the actual file before writing this parser):
  - A "Heading 1" paragraph "THỨ TỰ ƯU TIÊN ÔN TẬP" (priority-order summary table), then
    exactly 11 "Heading 1" paragraphs matching "CHỦ ĐỀ <n> — <name>" (one per topic, in
    document order = priority order), then a final "Heading 1" appendix
    "PHỤ LỤC A — TỪ ĐIỂN THUẬT NGỮ" (glossary; not parsed here, see glossary.py).
  - Within a topic, each question is three or more consecutive "Memory Card"-styled
    paragraphs: "<n>. <question text>", "Cốt lõi: <answer text>", and optionally
    "Hiểu sâu: <explanation text>" (a "Hiểu sâu" line is not always present -- some cards
    only carry a core answer, so `explanation` may be an empty string).
  - Question numbers are globally unique but NOT contiguous: the handbook is a
    priority-reordered subset of a larger original question bank, so ranges like 141-160 are
    intentionally absent. Table content (per-topic term tables, the priority-order table, the
    appendix) is automatically skipped because `Document.paragraphs` only ever yields
    body-level paragraphs, never paragraphs nested inside table cells.

Pure Python + python-docx only -- no FastAPI, no SQLAlchemy, importable standalone by both
scripts/extract_practical_review_docx.py and the FastAPI app.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as _open_docx

from app.practical_review.models import ParsedDocument, PracticalQuestion, PracticalTopic

# Positional (slug, ui_group) for the 11 topics, in the exact order they appear in the DOCX
# (which is the handbook's priority-study order). The display name itself is taken from the
# DOCX heading text, never hardcoded here -- slug/group are navigation metadata invented for
# this app, not source content.
_TOPIC_SLUGS_AND_GROUPS: list[tuple[str, str]] = [
    ("ai-coding", "Thực chiến"),
    ("du-an-production", "Thực chiến"),
    ("java-core", "Java"),
    ("oop-solid", "Java"),
    ("spring-boot", "Java"),
    ("git", "Backend nền tảng"),
    ("thuat-toan", "Backend nền tảng"),
    ("rest-api", "Backend nền tảng"),
    ("sql-database", "Backend nền tảng"),
    ("python-core", "Python"),
    ("python-backend", "Python"),
]

_TOPIC_HEADING_RE = re.compile(r"^CHỦ ĐỀ\s+\d+\s*—\s*(.*)$")
_QUESTION_RE = re.compile(r"^(\d+)\.\s*(.*)$")
_CORE_RE = re.compile(r"^Cốt lõi:\s*(.*)$")
_DEEP_RE = re.compile(r"^Hiểu sâu:\s*(.*)$")


class DocxStructureError(ValueError):
    """Raised when the DOCX doesn't match the expected practical-review structure. Carries
    enough detail to fix the source file or the parser, never silently drops data."""


@dataclass
class _PendingQuestion:
    number: int
    topic_slug: str
    topic_name: str
    question: str
    core: str | None = None
    deep: str | None = None


def parse_docx(path: str | Path) -> ParsedDocument:
    path = Path(path)
    if not path.exists():
        raise DocxStructureError(f"Không tìm thấy file DOCX: {path}")

    document = _open_docx(str(path))

    topics: list[PracticalTopic] = []
    questions: list[PracticalQuestion] = []
    topic_question_counts: dict[str, int] = {}
    seen_numbers: set[int] = set()

    current_slug: str | None = None
    current_topic_name: str | None = None
    current: _PendingQuestion | None = None

    def finalize_current() -> None:
        nonlocal current
        if current is None:
            return
        question_text = current.question.strip()
        core_text = current.core.strip() if current.core is not None else ""
        if not question_text or not core_text:
            raise DocxStructureError(
                f"Câu {current.number} thiếu nội dung câu hỏi hoặc phần 'Cốt lõi'."
            )
        deep_text = current.deep.strip() if current.deep is not None else ""
        questions.append(
            PracticalQuestion(
                number=current.number,
                topic_slug=current.topic_slug,
                topic_name=current.topic_name,
                question=question_text,
                answer=core_text,
                explanation=deep_text,
            )
        )
        topic_question_counts[current.topic_slug] += 1
        current = None

    for para in document.paragraphs:
        style_name = para.style.name if para.style else None
        text = para.text.strip()

        if style_name == "Heading 1":
            finalize_current()
            topic_match = _TOPIC_HEADING_RE.match(text)
            if topic_match is None:
                current_slug = None
                current_topic_name = None
                continue
            order = len(topics) + 1
            if order > len(_TOPIC_SLUGS_AND_GROUPS):
                raise DocxStructureError(
                    f"Tìm thấy nhiều hơn {len(_TOPIC_SLUGS_AND_GROUPS)} chủ đề "
                    f"(Heading 1 dạng 'CHỦ ĐỀ n — ...') trong DOCX; chủ đề thừa: {text!r}"
                )
            slug, group = _TOPIC_SLUGS_AND_GROUPS[order - 1]
            current_slug = slug
            current_topic_name = topic_match.group(1).strip()
            topic_question_counts[slug] = 0
            topics.append(
                PracticalTopic(
                    slug=slug,
                    order=order,
                    heading=text,
                    display_name=current_topic_name,
                    group=group,
                    question_count=0,  # filled in after the full pass below
                )
            )
            continue

        if style_name != "Memory Card" or not text:
            continue

        question_match = _QUESTION_RE.match(text)
        if question_match is not None:
            finalize_current()
            if current_slug is None or current_topic_name is None:
                raise DocxStructureError(
                    f"Tìm thấy thẻ câu hỏi {text!r} trước khi có Heading 1 chủ đề nào -- DOCX "
                    "không đúng cấu trúc mong đợi."
                )
            number = int(question_match.group(1))
            if number in seen_numbers:
                raise DocxStructureError(f"Số thứ tự câu hỏi bị trùng: {number}")
            seen_numbers.add(number)
            current = _PendingQuestion(
                number=number,
                topic_slug=current_slug,
                topic_name=current_topic_name,
                question=question_match.group(2),
            )
            continue

        if current is None:
            continue  # stray "Memory Card" text outside any question -- ignore defensively

        core_match = _CORE_RE.match(text)
        if core_match is not None:
            current.core = core_match.group(1)
            continue
        deep_match = _DEEP_RE.match(text)
        if deep_match is not None:
            current.deep = deep_match.group(1)

    finalize_current()

    if len(topics) != len(_TOPIC_SLUGS_AND_GROUPS):
        raise DocxStructureError(
            f"Tìm thấy {len(topics)} chủ đề (Heading 1 dạng 'CHỦ ĐỀ n — ...'), "
            f"cần đúng {len(_TOPIC_SLUGS_AND_GROUPS)}."
        )

    topics = [
        PracticalTopic(
            slug=t.slug,
            order=t.order,
            heading=t.heading,
            display_name=t.display_name,
            group=t.group,
            question_count=topic_question_counts.get(t.slug, 0),
        )
        for t in topics
    ]

    return ParsedDocument(topics=topics, questions=questions)
